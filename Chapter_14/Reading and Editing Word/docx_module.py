import docx

d = docx.Document('C:\\Emil\\Proiecte\\Python\\Proiecte_Python\\Automate_the_boring_stuff\\Chapter_14\\Reading and Editing Word\\demo.docx')

print(d.paragraphs)
print(d.paragraphs[0].text)  # print the first paragraph

p = d.paragraphs[1]
print(p.runs)  # a run is counted whenever a change in style happens to the current paragraph

print(p.runs[0].text)
print(p.runs[1].text)
print(p.runs[2].text)
print(p.runs[3].text)

print(p.runs[1].bold)  # print true if the style is applied to the paragraph

p.runs[3].underline = True  # change style to underline
p.runs[3].text = 'italic and underlined.'  # text changed
p.style = 'Title'
d.save('demo1.docx')  #  save the new PDF file

d = docx.Document()  # make a new blank document
d.add_paragraph('Hello this is a paragraph')
d.add_paragraph('This is another paragraph')
d.save('demo2.docx')

p = d.paragraphs[0]
p.add_run('This is a new run')
p.runs[1].bold = True
d.save('demo3.docx')